"""Author the enrolment authorisation Doctavian renders and a person signs.

This is the document the whole Foxit argument turns on. It states, in words a
non engineer can hold someone to, exactly what publishing a key to a domain
means and how long it lasts. Somebody puts their name to that claim before any
key becomes live.

Two things are load bearing in the markup.

The authorisation hash is printed as ordinary body text. The broker downloads
the executed document, reads it back as text, and looks for that string. A
webhook field says a person acted; the hash says what they acted on, and only
the second is something we put there ourselves.

The signature and date fields are eSign text tags rather than coordinates. A
tag travels with the sentence it belongs to, so a longer diligence section moves
the signature block with it. Coordinates guessed at authoring time do not move,
and a signature box that lands on the wrong paragraph is worse than one that is
hard to place.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUT = Path("assets/signet-authorisation.docx")


def line(
    document: Document,
    text: str,
    *,
    size: int = 10,
    bold: bool = False,
    space_after: int = 4,
    mono: bool = False,
    invisible: bool = False,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(space_after)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Courier New" if mono else "Helvetica"
    if invisible:
        # White on white. Both vendors recommend exactly this: their markers
        # stay in the file where the platform can find them, and the person
        # being asked to sign never sees the machinery.
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def build() -> None:
    document = Document()
    for section in document.sections:
        section.left_margin = section.right_margin = Pt(54)
        section.top_margin = section.bottom_margin = Pt(54)

    line(document, "SIGNET", size=20, bold=True, space_after=0)
    line(document, "Enrolment authorisation", size=9, space_after=18)

    line(
        document,
        "{!Enrolment.Brand} at {!Enrolment.Domain}",
        size=14,
        bold=True,
        space_after=10,
    )
    line(
        document,
        "This authorises Signet to publish a signing key in the DNS of "
        "{!Enrolment.Domain}. Read the next paragraph before signing.",
        space_after=14,
    )

    line(document, "WHAT YOU ARE AGREEING TO", size=9, bold=True, space_after=4)
    line(
        document,
        "Once this record is live, any document signed with the matching private "
        "key will verify as issued by {!Enrolment.Domain}, to anyone, anywhere, "
        "with no further involvement from you. There is no expiry. Removing the "
        "DNS record is the only way to stop it, and documents already signed "
        "keep verifying until you do.",
        space_after=10,
    )
    line(
        document,
        "Sign this only if you control the DNS for {!Enrolment.Domain} and intend "
        "that domain to vouch for documents issued in its name.",
        space_after=14,
    )

    line(document, "THE RECORD TO BE PUBLISHED", size=9, bold=True, space_after=4)
    line(document, "_signet.{!Enrolment.Domain}   TXT", size=9, mono=True, space_after=2)
    line(document, "{!Enrolment.Record}", size=8, mono=True, space_after=10)
    line(document, "Key fingerprint   {!Enrolment.Fingerprint}", size=9, mono=True, space_after=14)

    line(document, "WHAT WAS CHECKED BEFORE ASKING", size=9, bold=True, space_after=4)
    line(document, "{!Enrolment.Diligence}", size=9, space_after=14)

    # Enrolment starts from whatever the issuer actually sent: a forwarded
    # thread, a chat log, a scan. None of it is labelled, and the fields above
    # were read out of it. Printing each field beside the line it came from is
    # what makes this document worth signing rather than worth clicking through.
    line(document, "HOW THIS WAS READ", size=9, bold=True, space_after=4)
    line(
        document,
        "Each field above was read out of the request as it arrived. The line it "
        "came from is printed beside it. Correct anything that is wrong before signing.",
        size=9,
        space_after=8,
    )
    line(
        document,
        '<mdoc:repeater name="readings" value="{!Enrolment.Readings}" '
        'variable="item" mode="standard">',
        space_after=2,
    )
    line(document, "{!#item#.Field}   {!#item#.Value}", size=9, bold=True, space_after=1)
    line(
        document,
        '   read from "{!#item#.Quote}"   ·   {!#item#.Note}',
        size=8,
        space_after=6,
    )
    line(document, '</mdoc:repeater name="readings">', space_after=8)

    # Branches on the readings themselves. A request that said one thing plainly
    # gets no warning; one the text supported two answers for gets a warning
    # naming that, decided here rather than by the caller.
    line(
        document,
        '<mdoc:paragraph name="ambiguous" '
        'hidden="{!$toDecimal(sum(Enrolment.Readings, "Uncertain")) < 1}">',
        space_after=2,
    )
    line(
        document,
        "At least one field above had more than one possible reading in the text. "
        "Those lines name the reading that was rejected. Check them against what "
        "you meant before putting your name to this.",
        size=9,
        bold=True,
        space_after=2,
    )
    line(document, '</mdoc:paragraph name="ambiguous">', space_after=14)

    # The string the broker looks for in the executed document. Printed rather
    # than encoded, so a person can compare it against what they were sent.
    line(document, "AUTHORISATION REFERENCE", size=9, bold=True, space_after=4)
    line(document, "{!Enrolment.AuthorisationHash}", size=9, mono=True, space_after=2)
    line(
        document,
        "Quote this reference in any query. The key is published only after a "
        "signed copy of this document carrying this exact reference is read back.",
        size=8,
        space_after=20,
    )

    # Two signing surfaces on one page, because the human gate is a port with
    # two implementations and either can be the one that runs.
    #
    # Foxit reads text tags: ${signfield:1:y}. Doctavian finds an anchor string
    # and places its field over the bounding box, leaving the text in place.
    # Both are inert marks to the other, so a document carrying both renders
    # correctly whichever gateway sends it.
    # Two signing surfaces on one page, because the human gate is a port with
    # two implementations and either can be the one that runs. Foxit reads text
    # tags; Doctavian lays its field over the bounding box of an anchor string.
    # Each is inert to the other.
    #
    # Both are printed white on white, which is what both vendors recommend,
    # and each sits on its own line under its label. The first version put them
    # inline, and the signing view showed a person raw ${signfield:1:y} markup
    # with a field dropped on top of it. A document somebody is being asked to
    # read and sign cannot look like a bug.
    #
    # Anchor length matters: the field is sized from the anchor's bounding box,
    # so a short anchor gives a signature box too small to sign in.
    line(document, "SIGNED FOR {!Enrolment.Brand}", size=9, bold=True, space_after=10)

    line(document, "Name", size=9, bold=True, space_after=2)
    line(document, "${textfield:1:y:Signer_Name:________________}", space_after=1, invisible=True)
    line(document, "_" * 46, space_after=10)

    line(document, "Role", size=9, bold=True, space_after=2)
    line(document, "${textfield:1:y:Signer_Role:________________}", space_after=1, invisible=True)
    line(document, "_" * 46, space_after=10)

    line(document, "Signature", size=9, bold=True, space_after=2)
    # Foxit takes explicit pixel dimensions after the field name. Doctavian
    # takes the anchor's bounding box, so that line is set large: a marker at
    # body size produces a box too small to sign in, which is what the first
    # envelope looked like.
    line(document, "${signfield:1:y:signature_issuer:220:64}", space_after=1, invisible=True)
    line(document, "_SIGNET_SIGNATURE_ISSUER_HERE_", size=20, space_after=6, invisible=True)
    line(document, "_" * 46, space_after=10)

    line(document, "Date", size=9, bold=True, space_after=2)
    line(document, "${datefield:1:y:Date_Signed:130:30}", space_after=1, invisible=True)
    line(document, "_SIGNET_DATE_ISSUER_HERE_", size=13, space_after=2, invisible=True)
    line(document, "_" * 30, space_after=0)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUT)
    print(f"  wrote {OUT} ({OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    build()
    sys.exit(0)
