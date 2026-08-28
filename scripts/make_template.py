"""Author the Signet invoice template Doctavian renders from.

Their markup is plain text, not Word content controls, which is what makes a
template authorable from a script rather than by hand in Word. Verified by
reading their own mission-1-agreement.docx: no w:sdt anywhere, repeaters and
conditional paragraphs written as mdoc tags in the body text.

Only expressions observed in that file are used here. addDays looks like it
ought to exist and is not written anywhere, so the due date is passed as data
instead of derived: a template that fails to render is worth less than one that
asks its caller for one more field.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUT = Path("assets/northpost-invoice.docx")

MONEY = "'number', '#,###.00'"
TOTAL = f'{{!$format(sum(Invoice.LineItems, "LineAmount"), {MONEY})}}'


def line(
    document: Document,
    text: str,
    *,
    size: int = 10,
    bold: bool = False,
    align: str = "left",
    space_after: int = 4,
    mono: bool = False,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
    }[align]
    paragraph.paragraph_format.space_after = Pt(space_after)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Courier New" if mono else "Helvetica"


def build() -> None:
    document = Document()
    for section in document.sections:
        section.left_margin = section.right_margin = Pt(54)
        section.top_margin = section.bottom_margin = Pt(54)

    line(document, "{!Invoice.PayeeMark}", size=20, bold=True, space_after=0)
    line(
        document,
        "{!Invoice.PayeeName}  ·  {!Invoice.PayeeDomain}",
        size=9,
        space_after=18,
    )

    line(document, "INVOICE {!Invoice.Number}", size=14, bold=True, space_after=2)
    line(document, "Billed to {!Invoice.CustomerName}", space_after=2)
    line(
        document,
        "Issued {!$format(date(Invoice.IssuedOn), 'date', 'medium')}"
        "   ·   Payable by {!$format(date(Invoice.DueOn), 'date', 'medium')}",
        space_after=16,
    )

    line(document, "SERVICES", size=9, bold=True, space_after=6)

    # Loops. One template, however many lines arrive.
    line(
        document,
        '<mdoc:repeater name="lines" value="{!Invoice.LineItems}" variable="item" mode="standard">',
        space_after=2,
    )
    line(
        document,
        "{!#item#.Description}    "
        "{!#item#.Quantity} x " + f"{{!$format(toDecimal(#item#.UnitPrice), {MONEY})}}"
        "    " + f"{{!$format(toDecimal(#item#.LineAmount), {MONEY})}}",
        space_after=2,
    )
    line(document, '</mdoc:repeater name="lines">', space_after=10)

    # Calculates. The total is never sent, it is derived from the lines.
    line(
        document,
        f"TOTAL DUE   {{!Invoice.Currency}} {TOTAL}",
        size=13,
        bold=True,
        align="right",
        space_after=16,
    )

    line(document, "PAY TO", size=9, bold=True, space_after=4)
    line(document, "Account name   {!Invoice.PayeeName}", mono=True, space_after=2)
    line(document, "IBAN           {!Invoice.Iban}", mono=True, space_after=2)
    line(document, "BIC            {!Invoice.Bic}", mono=True, space_after=14)

    # Branches. The clause that matters only above a threshold appears only
    # above that threshold, decided by the template rather than by the caller.
    line(
        document,
        '<mdoc:paragraph name="large" '
        'hidden="{!$toDecimal(sum(Invoice.LineItems, "LineAmount")) < 10000}">',
        space_after=2,
    )
    line(
        document,
        "This invoice exceeds 10,000. Confirm the account above against a "
        "previously trusted record before transferring. Northpost never changes "
        "bank details by email.",
        size=9,
        space_after=2,
    )
    line(document, '</mdoc:paragraph name="large">', space_after=14)

    # The signature travels in the code, not as body text. Printing the signed
    # fields in words put iban= and amt= on the page in a form an extractor
    # reads as page data, and a document that prints the answer to the question
    # being asked about it can agree with itself while the payment block says
    # something else. Measured: a doctored invoice came back certified because
    # extraction took the account number off the printed mark instead of the
    # PAY TO block. Scanning the code gives the same string to anyone who wants
    # to check it by hand.
    line(document, "PROOF OF ORIGIN", size=9, bold=True, space_after=4)
    line(document, "Verify at {!Invoice.SignetLocator}", size=9, space_after=2)
    line(
        document,
        "Scan the code for the signature. The key is published at "
        "_signet.{!Invoice.PayeeDomain}. Check it with dig and openssl, without an "
        "account here.",
        size=8,
        space_after=26,
    )

    # A real invoice fills its page. The terms below are ordinary for freight
    # and they are also what an interception attempt has to talk its way past,
    # which is why the remittance reference is stated rather than assumed.
    line(document, "PAYMENT TERMS", size=9, bold=True, space_after=4)
    line(
        document,
        "Net 30 days from the date of issue. Interest on overdue amounts accrues "
        "at 8 percent above base rate from the day after the due date.",
        size=9,
        space_after=10,
    )
    line(document, "REMITTANCE", size=9, bold=True, space_after=4)
    line(
        document,
        "Quote {!Invoice.Number} as the payment reference. Send remittance advice "
        "to accounts@{!Invoice.PayeeDomain}. Payments without a reference are held "
        "unallocated.",
        size=9,
        space_after=10,
    )
    line(document, "QUERIES", size=9, bold=True, space_after=4)
    line(
        document,
        "Raise any query within 14 days of the invoice date, quoting the invoice "
        "number. Bank details are never amended by email or telephone. A request "
        "to pay a different account is a fraud attempt and should be reported.",
        size=9,
        space_after=10,
    )
    line(
        document,
        "{!Invoice.PayeeName}  ·  {!Invoice.PayeeDomain}  ·  page 1 of 1",
        size=8,
        space_after=0,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUT)
    print(f"  wrote {OUT} ({OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    build()
    sys.exit(0)
